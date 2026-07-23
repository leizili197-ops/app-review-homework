#!/usr/bin/env python3
"""App Store Review Analyzer — 美区 App Store 评论抓取 + DeepSeek AI 分析 + PRD/测试用例生成

历史修复：
  Bug 1：链接解析不健壮 → 用正则 r'id(\\d+)' 提取 App ID，兼容多种格式，国家码非 us 报错
  Bug 2：生成的 .docx 打不开 → 用 python-docx 生成真·Word 文档，并生成后立刻自检

本版本优化：
  - 评论抓取同时支持 Apple RSS 的 **XML 格式**（/xml，xml.etree.ElementTree 解析）
    和 **JSON 格式**（/json）两种接口
  - 实测发现：Apple 对"某个 App 能用哪个格式返回数据"没有规律 —— 有的 App 只有 XML 有、
    有的只有 JSON 有、有的都没有（Apple 接口限制）。因此统一入口改成「XML/JSON 都试，
    取拿到（数量多）的那一份」，最大化拿到真实评论的概率
  - 新增 compare_json_vs_xml()：终端对比 JSON / XML 两种方式抓取条数，方便确认稳定性
  - 抓取为 0 条时不再崩溃，给出中文说明 + 手动粘贴评论的替代入口
"""

# ============================================================================
# 1. 导入需要的库
# ============================================================================
import streamlit as st          # Streamlit：搭网页界面
import json                     # 处理 JSON（评论缓存、DeepSeek 返回结果）
import re                       # 正则表达式：从链接里抠 App ID、从文本里抠评分
import io                       # io.BytesIO：内存里存二进制数据（生成 docx 下载用）
import time                     # time.sleep：抓取时稍微停顿，避免请求过快被限制
import hashlib                  # 根据评论内容生成唯一短 ID
import xml.etree.ElementTree as ET  # 标准库：解析 Apple RSS 返回的 XML 评论
from datetime import datetime  # 记录抓取/分析时间
from pathlib import Path        # 处理文件路径
from typing import Optional     # 类型标注：表示"可能有值，也可能没有"

import requests                 # 发起 HTTP 请求（抓评论、调 DeepSeek）
from docx import Document       # python-docx：生成 Word(.docx) 文档的核心库
from docx.shared import Pt, RGBColor  # 控制字体大小、颜色（可选美化）

# 说明：app_store_scraper 采用「按需懒加载」——只在真正用到它做备用抓取时才 import。
# 原因：该库上游把 requests 死锁在 2.23.0，与 streamlit 需要的 >=2.27 冲突，
# 如果写在文件顶部强制导入，一旦没装这个库，整个应用都会启动失败。
# 改成懒加载后，即使没装它，应用也能靠主路径（Apple RSS）正常运行；
# 想启用备用抓取时，单独执行：pip install --no-deps app-store-scraper 即可。


# ============================================================================
# 2. 页面配置 & 常量
# ============================================================================
st.set_page_config(
    page_title="App Review Analyzer",
    page_icon="📱",
    layout="wide",                 # 宽屏布局
    initial_sidebar_state="expanded",
)

CACHE_FILE = Path(__file__).parent / "reviews_cache.json"   # 评论缓存文件名
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

# Word 文档正确的 MIME 类型（下载按钮必须填这个，浏览器才会当 Word 文件处理）
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 记录最近一次实时抓取实际用到的数据源（"XML" / "JSON" / ""）
# 用模块级全局变量，UI 展示时读取；裸跑（非 Streamlit 环境）也不会报错。
LAST_FETCH_SOURCE = ""


# ============================================================================
# 3. Bug 1 修复：健壮的链接解析
# ============================================================================
def parse_apple_store_url(url: str) -> tuple[Optional[str], Optional[str]]:
    """从任意格式的美区 App Store 链接里提取出数字 App ID。

    【返回值设计】
      返回 (app_id, error_message) 这样一个"两个元素的元组"：
      - 成功时：返回 (AppID字符串, None)         例如 ("835599320", None)
      - 失败时：返回 (None, 中文错误提示)         例如 (None, "未识别到有效的美区 App ID...")

    【为什么这么设计】
      把"结果"和"错误原因"一起返回，调用方拿到后直接判断：
        如果 error 不是 None → 用 st.error() 把错误显示给用户；
        否则就用 app_id 去抓取。这样程序永远不会"静默失败"或崩溃。

    【支持格式举例】（都用正则 r'id(\\d+)' 匹配，不依赖链接分段位置）
      1. https://apps.apple.com/us/app/应用名/id123456789
      2. https://apps.apple.com/us/app/id123456789
      3. https://itunes.apple.com/us/app/应用名/id123456789?mt=8
      4. 末尾带斜杠 / 或带其他查询参数 ?l=zh-Hans-CN 的情况
    """
    if not url or not url.strip():
        return None, "请输入 App Store 链接"

    url = url.strip()  # 去掉首尾可能存在的空格/换行

    # ---------- 第一步：确认这是"美区"(us) 链接 ----------
    # 美区链接的路径里一定包含 "/us/"，比如 /us/app/...
    # 用正则抓路径里第一个 "/两个小写字母/" 片段，例如 /us/、/cn/、/jp/
    country_match = re.search(r"/([a-z]{2})/", url.lower())
    if not country_match:
        return None, "未识别到有效的国家码，请确认是美区链接（链接中应包含 /us/）"
    if country_match.group(1) != "us":
        return None, (
            f"检测到国家码为「{country_match.group(1)}」，"
            f"本工具只支持美区（us）App Store，请使用 /us/ 链接"
        )

    # ---------- 第二步：用正则提取纯数字 App ID ----------
    # 关键修复点：直接在整个链接里搜索 "id" + 一串数字，
    # 不关心 id 出现在第几段、前后是什么文字，非常健壮。
    id_match = re.search(r"id(\d+)", url)
    if not id_match:
        return None, "未识别到有效的美区 App ID，请检查链接格式（应包含如 /id123456789 的数字 ID）"

    app_id = id_match.group(1)  # group(1) 就是括号里捕获到的纯数字
    return app_id, None


def extract_app_name(url: str) -> str:
    """从链接里尽量取出 App 名称（仅用于界面展示，失败也不影响主流程）。"""
    m = re.search(r"/app/([^/]+)/", url)
    if m:
        return m.group(1).replace("-", " ").title()
    return "Unknown App"


def run_url_parser_tests() -> list:
    """链接解析"自测"：用一组典型链接验证解析是否正确。

    返回格式：[(链接, 期望ID, 实际ID或错误, 是否通过), ...]
    界面上的"🧪 链接解析自测"按钮会调用它并把结果画成表格。
    """
    test_cases = [
        ("https://apps.apple.com/us/app/tiktok/id835599320", "835599320"),
        ("https://apps.apple.com/us/app/id123456789", "123456789"),
        ("https://itunes.apple.com/us/app/spotify/id324684580?mt=8", "324684580"),
        ("https://apps.apple.com/us/app/some-app/id987654321/", "987654321"),
        ("https://apps.apple.com/us/app/app-name/id111222333?l=zh-Hans-CN", "111222333"),
        ("https://apps.apple.com/cn/app/app-name/id445678123", "cn-should-fail"),
        ("https://example.com/abcdefg", "no-id-should-fail"),
    ]
    results = []
    for url, expected in test_cases:
        app_id, err = parse_apple_store_url(url)
        if err:
            passed = expected.endswith("should-fail")
            results.append((url, "报错(符合预期)" if passed else "报错(不符合预期)", err, passed))
        else:
            passed = (app_id == expected)
            results.append((url, expected, app_id, passed))
    return results


# ============================================================================
# 4. 优化：评论抓取（Apple RSS — 主用 XML，回退 JSON）
# ============================================================================
def _localname(tag: str) -> str:
    """把一个 XML 标签名（可能带命名空间）取出"本地名"。

    例如：
      '{http://www.w3.org/2005/Atom}entry' -> 'entry'
      '{http://itunes.apple.com/rss}rating' -> 'rating'   （也兼容 'im:rating'）
    这样解析 Apple 的 RSS XML 时就不用死记硬背那一长串命名空间 URL。
    """
    return tag.split("}")[-1].split(":")[-1]


def parse_apple_rss_xml(xml_text: str) -> list:
    """解析 Apple RSS 评论的 **XML** 文本，返回评论字典列表。

    每条评论包含：评论ID、标题、内容、评分、版本号（外加作者、日期，便于展示）。

    Apple 的评论 XML 结构（简化）大致是：
      <feed ...>
        <entry>                       <!-- 第 1 个 entry 是 App 自身信息，无评分，会被过滤掉 -->
          <id>https://itunes.apple.com/review?id=xxxx</id>
          <title>评论标题</title>
          <content type="text">评论正文</content>
          <im:rating>5</im:rating>
          <im:version>12.3.0</im:version>
          <author><name>用户名</name></author>
          <updated>2024-01-01T00:00:00-07:00</updated>
        </entry>
        ...
      </feed>
    """
    reviews: list = []
    try:
        root = ET.fromstring(xml_text)   # 把 XML 字符串解析成元素树
    except ET.ParseError:
        return reviews                    # 解析失败就返回空列表，不崩溃

    # 遍历整棵树，找出所有 <entry> 元素（忽略命名空间）
    entries = [e for e in root.iter() if _localname(e.tag) == "entry"]

    for entry in entries:
        # 把 entry 下的子元素按"本地名"建个快速查找字典：{'title': 元素, 'rating': 元素, ...}
        fields = {_localname(c.tag): c for c in entry}

        # 过滤掉"App 自身信息"那条 entry：真实的用户评论一定带有 <rating>
        if "rating" not in fields:
            continue

        # ---- 评论ID：从 <id> 文本里抠数字（Apple 给的是 review?id=xxxx 这种链接）----
        id_text = (fields.get("id").text or "") if fields.get("id") is not None else ""
        m = re.search(r"id=(\d+)", id_text)
        rid = m.group(1) if m else hashlib.md5(id_text.encode()).hexdigest()[:8]

        # ---- 标题 ----
        title_el = fields.get("title")
        title = (title_el.text or "").strip() if title_el is not None else ""

        # ---- 内容（评论正文）----
        content_el = fields.get("content")
        content = (content_el.text or "").strip() if content_el is not None else ""

        # ---- 评分（同时做范围校验：只保留 1~5 的有效评分，进一步过滤异常条目）----
        rating_el = fields.get("rating")
        try:
            rating = int((rating_el.text or "0").strip()) if rating_el is not None else 0
        except (ValueError, TypeError):
            rating = 0
        if rating < 1 or rating > 5:
            continue

        # ---- 版本号 ----
        version_el = fields.get("version")
        version = (version_el.text or "").strip() if version_el is not None else ""

        # ---- 作者名（<author><name>用户名</name></author>）----
        user_name = "Anonymous"
        author_el = fields.get("author")
        if author_el is not None:
            for child in author_el:
                if _localname(child.tag) == "name":
                    user_name = (child.text or "Anonymous").strip()
                    break

        # ---- 日期 ----
        updated_el = fields.get("updated")
        date = (updated_el.text or "").strip() if updated_el is not None else ""

        reviews.append({
            "id": rid,
            "userName": user_name,
            "title": title,
            "review": content,
            "rating": rating,
            "date": date,
            "version": version,
        })

    return reviews


def _fetch_rss_xml(app_id: str, country: str = "us", max_count: int = 200) -> list:
    """用 **XML 格式** 抓取 Apple RSS 评论（主路径）。

    接口地址示例：
      https://itunes.apple.com/us/rss/customerreviews/id=835599320/sortby=mostrecent/xml
    支持翻页（page=2,3,...）。
    """
    all_reviews: list = []
    seen: set = set()

    for page in range(1, 11):  # 最多翻 10 页
        if len(all_reviews) >= max_count:
            break

        # 第一页和后续分页的 URL 略有不同（page 段放在 id 之前）
        if page == 1:
            url = f"https://itunes.apple.com/{country}/rss/customerreviews/id={app_id}/sortby=mostrecent/xml"
        else:
            url = f"https://itunes.apple.com/{country}/rss/customerreviews/page={page}/id={app_id}/sortby=mostrecent/xml"

        try:
            resp = requests.get(url, timeout=30)
            if resp.status_code != 200:
                break

            entries = parse_apple_rss_xml(resp.text)  # 解析这一页的 XML
            if not entries:
                break  # 没有评论了，结束翻页

            for e in entries:
                if e["id"] in seen:
                    continue
                seen.add(e["id"])
                all_reviews.append(e)

            if len(entries) < 50:  # 已经到最后一页（每页最多 50 条）
                break

            time.sleep(0.3)  # 礼貌地停顿，避免请求过快

        except Exception:
            break  # 任何异常都安全退出，不向上抛

    return all_reviews


def _fetch_rss_json(app_id: str, country: str = "us", max_count: int = 200) -> list:
    """用 **JSON 格式** 抓取 Apple RSS 评论（仅作为 XML 失败时的回退）。"""
    all_reviews: list = []
    seen: set = set()

    for page in range(1, 11):
        if len(all_reviews) >= max_count:
            break

        if page == 1:
            url = f"https://itunes.apple.com/{country}/rss/customerreviews/id={app_id}/sortBy=mostRecent/json"
        else:
            url = f"https://itunes.apple.com/{country}/rss/customerreviews/page={page}/id={app_id}/sortby=mostrecent/json"

        try:
            resp = requests.get(url, timeout=30)
            if resp.status_code != 200:
                break

            data = resp.json()
            entries = data.get("feed", {}).get("entry", [])
            if not entries:
                break  # 这一页真的没有评论了

            for entry in entries:
                # 只保留"带有效评分(1~5)"的条目：
                # Apple 有时把 App 自身信息放在第一条（无评分），有时第一条就是真实评论，
                # 用"评分必须在 1~5 之间"来过滤，比"跳过第一条"更稳，不受格式变化影响。
                try:
                    rating = int(entry.get("im:rating", {}).get("label", "0"))
                except (ValueError, TypeError):
                    rating = 0
                if rating < 1 or rating > 5:
                    continue

                rid = entry.get("id", {}).get("label", "")
                if rid in seen:
                    continue
                seen.add(rid)
                all_reviews.append({
                    "id": hashlib.md5(
                        f"{entry.get('author', {}).get('name', {}).get('label', '')}"
                        f"{entry.get('updated', {}).get('label', '')}"
                        f"{entry.get('content', {}).get('label', '')}".encode()
                    ).hexdigest()[:8],
                    "userName": entry.get("author", {}).get("name", {}).get("label", "Anonymous"),
                    "title": entry.get("title", {}).get("label", ""),
                    "review": entry.get("content", {}).get("label", ""),
                    "rating": rating,
                    "date": entry.get("updated", {}).get("label", ""),
                    "version": entry.get("im:version", {}).get("label", ""),
                })

            if len(entries) < 50:
                break
            time.sleep(0.3)
        except Exception:
            break

    return all_reviews


def fetch_reviews_via_rss(app_id: str, country: str = "us", max_count: int = 200) -> list:
    """抓取评论的统一入口：XML 与 JSON 都试，取拿到（数量多）的那一份。

    为什么要"两种都试"而不是"XML 优先、JSON 兜底"？
    实测多个 App 后发现，Apple 的 RSS 接口对"哪个格式能返回数据"没有规律：

      - Spotify：XML 有 50 条，JSON 一条都没有  -> 若只用 JSON 就彻底失败
      - TikTok  ：JSON 有 50 条，XML 一条都没有  -> 若只用 XML 就浪费一次请求
      - Netflix ：两种都有（XML 50 / JSON 49）
      - Instagram/YouTube/WhatsApp：两种都为 0   -> 这是 Apple 接口限制，非代码 bug

    所以最稳的做法是两边都抓、取有数据且数量多的那一份；都空才返回 []，
    交给上层去提示"手动粘贴评论"。
    """
    global LAST_FETCH_SOURCE

    xml_reviews = _fetch_rss_xml(app_id, country, max_count)
    json_reviews = _fetch_rss_json(app_id, country, max_count)

    # 两边都空 -> 标记来源为空，返回 []（上层据此触发"手动粘贴"）
    if not xml_reviews and not json_reviews:
        LAST_FETCH_SOURCE = ""
        return []

    # 数量相同时优先 XML（贴合"以 XML 为主"的需求偏好）；否则取多的那一方
    if len(xml_reviews) >= len(json_reviews):
        source, reviews = ("XML", xml_reviews) if xml_reviews else ("JSON", json_reviews)
    else:
        source, reviews = ("JSON", json_reviews)

    LAST_FETCH_SOURCE = source  # 记下真实用到的来源，供 UI 展示（裸跑也不报错）
    return reviews


def compare_json_vs_xml(app_id: str, country: str = "us") -> tuple:
    """对比同一 App 分别用 JSON 和 XML 两种方式抓取到的评论条数（终端打印）。

    用法（在终端里执行）：
        python3 -c "import app; app.compare_json_vs_xml('835599320')"
    """
    # ---- JSON 方式 ----
    json_count = 0
    try:
        url = f"https://itunes.apple.com/{country}/rss/customerreviews/id={app_id}/sortBy=mostRecent/json"
        data = requests.get(url, timeout=30).json()
        entries = data.get("feed", {}).get("entry", [])
        # 第一条是 App 自身信息，不算评论
        json_count = max(0, len(entries) - 1) if entries else 0
    except Exception as e:
        print(f"JSON 方式出错: {e}")

    # ---- XML 方式 ----
    xml_count = 0
    try:
        url = f"https://itunes.apple.com/{country}/rss/customerreviews/id={app_id}/sortby=mostrecent/xml"
        resp = requests.get(url, timeout=30)
        xml_count = len(parse_apple_rss_xml(resp.text))
    except Exception as e:
        print(f"XML 方式出错: {e}")

    print(f"App ID {app_id} 抓取对比：")
    print(f"  JSON 方式抓到: {json_count} 条")
    print(f"  XML  方式抓到: {xml_count} 条")
    if xml_count > json_count:
        print(f"  => XML 比 JSON 多 {xml_count - json_count} 条（XML 更稳定）")
    elif xml_count < json_count:
        print(f"  => JSON 比 XML 多 {json_count - xml_count} 条")
    else:
        print("  => 两种方式数量一致")
    return json_count, xml_count


def parse_manual_reviews(raw_text: str) -> list:
    """把用户在文本框里手动粘贴的评论，解析成统一的评论字典列表。

    支持的粘贴格式（简单直观）：
      - 用「空一行」分隔两条评论
      - 每条里可用「评分: 数字」标注星级，不写评分默认按 3 星处理
      - 其余文字都当作评论内容

    示例：
      这个 App 老是闪退，太难受了
      评分: 1

      新版本好多了，推荐！
      评分: 5
    """
    reviews: list = []
    if not raw_text or not raw_text.strip():
        return reviews

    # 用「连续两个及以上换行」把文本切成多条评论块
    blocks = re.split(r"\n\s*\n", raw_text.strip())
    for idx, block in enumerate(blocks):
        block = block.strip()
        if not block:
            continue

        # 尝试从块里找出评分（如 "评分: 5" / "5星" / "rating: 4"）
        rating = 3  # 默认 3 星
        m = re.search(r"(?:评分|rating|星)[:：]?\s*(\d)", block, re.IGNORECASE)
        if m:
            rating = int(m.group(1))
            # 把"评分那一行"从内容里去掉，避免混进正文
            block = re.sub(
                r".*(?:评分|rating|星)[:：]?\s*\d.*\n?", "", block, flags=re.IGNORECASE
            ).strip()

        content = block
        reviews.append({
            "id": hashlib.md5(f"manual-{idx}-{content}".encode()).hexdigest()[:8],
            "userName": "手动录入",
            "title": content[:20],
            "review": content,
            "rating": rating,
            "date": datetime.now().strftime("%Y-%m-%d"),
            "version": "手动录入",
        })
    return reviews


# ============================================================================
# 5. 缓存读写
# ============================================================================
def load_cache() -> Optional[dict]:
    """读取本地缓存的评论数据（仅作可选离线参考，v2.0 不再作为自动兜底）。"""
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return None
    return None


def save_cache(reviews: list, app_name: str, app_id: str):
    """把抓到的评论写入 reviews_cache.json。"""
    cache_data = {
        "app_name": app_name,
        "app_id": app_id,
        "fetched_at": datetime.now().isoformat(),
        "total_reviews": len(reviews),
        "reviews": reviews,
    }
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache_data, f, ensure_ascii=False, indent=2)


# ============================================================================
# 6. 调用 DeepSeek 大模型做分析
# ============================================================================
def call_deepseek(prompt: str, api_key: str, max_tokens: int = 4096) -> str:
    """调用 DeepSeek 的对话接口，返回模型生成的文本。

    出错时会抛出 RuntimeError，附带按 HTTP 状态码分类的中文操作建议，
    让界面上的报错信息对用户更有用（而不是只显示一个英文错误码）。
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {
                "role": "system",
                "content": (
                    "你是一位资深的产品分析师，擅长分析用户评论、撰写产品需求文档（PRD）"
                    "和生成测试用例。你必须严格基于提供的评论数据进行分析，每条结论都要"
                    "引用具体的评论ID或原文。回复使用中文，结构清晰，使用 Markdown 格式。"
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens,
    }

    try:
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)
    except requests.exceptions.Timeout:
        raise RuntimeError(
            "DeepSeek API 超时（>120 秒）。可能原因：\n"
            "1）网络不稳定；\n"
            "2）当前请求的评论太多导致处理时间过长；\n"
            "建议：稍后重试，或在侧边栏把抓取数量调小一些再试。"
        )
    except requests.exceptions.ConnectionError as e:
        raise RuntimeError(
            f"无法连接到 DeepSeek 服务器（{e}）。请检查网络是否能访问外网。"
        )

    # ---- 按状态码给出不同的中文操作建议 ----
    if resp.status_code == 200:
        data = resp.json()
        return data["choices"][0]["message"]["content"]

    if resp.status_code == 401:
        raise RuntimeError(
            "DeepSeek API Key 无效（401 Unauthorized）。\n"
            "请检查：\n"
            "1）Key 是否复制完整（有没有多/少字符）；\n"
            "2）Key 是否已过期；\n"
            "获取新 Key：https://platform.deepseek.com"
        )
    if resp.status_code == 402:
        raise RuntimeError(
            "DeepSeek 余额不足（402 Insufficient Balance）。\n"
            "请前往 https://platform.deepseek.com 充值后再试。\n"
            "充值后无需重启应用，直接重新点「开始分析」即可。"
        )
    if resp.status_code == 429:
        raise RuntimeError(
            "DeepSeek 请求过于频繁（429 Rate Limited）。\n"
            "请等待 30~60 秒后重试。"
        )
    if resp.status_code >= 500:
        raise RuntimeError(
            f"DeepSeek 服务端暂时不可用（{resp.status_code}）。\n"
            "这是 DeepSeek 官方服务的问题，不是我们应用的 bug。\n"
            "建议：稍后重试。"
        )

    # 其他未知状态码
    try:
        detail = resp.json().get("error", {}).get("message", resp.text[:300])
    except Exception:
        detail = resp.text[:300]
    raise RuntimeError(f"DeepSeek API 返回异常（{resp.status_code}）：{detail}")


def build_analysis_prompt(reviews: list, user_goal: str) -> str:
    """把评论数据 + 用户目标拼成发给 DeepSeek 的提示词（prompt）。"""
    reviews_for_analysis = reviews[:200]  # 最多取前 200 条控制 token
    reviews_text = json.dumps(reviews_for_analysis, ensure_ascii=False, indent=2)

    prompt = f"""请分析以下 App Store 用户评论数据，完成以下三个任务。

用户特别关注的方向：{user_goal}

评论数据（JSON 格式，每条包含 id、userName、title、review、rating、date、version）：
```json
{reviews_text}
```

---

### 任务一：评论核心问题分析
请深入分析评论内容，找出用户抱怨的核心问题。要求：
1. **不能只靠关键词匹配**，必须理解语义和上下文
2. 每条结论必须**引用至少 2 条具体评论的 ID 和原文**作为证据
3. 按问题严重程度排序（抱怨人数从多到少）
4. 对每个问题给出影响评估（影响用户数估算、严重程度：高/中/低）
5. 将问题归类为：Bug类、体验类、付费/订阅类、功能缺失类、性能类、内容类等

### 任务二：产品需求文档（PRD）
基于上述问题生成 PRD。要求：
1. 每个需求必须**明确关联到任务一中的具体问题**
2. 按优先级划分 P0 / P1 / P2
3. 按版本拆分 V1.0（紧急修复）/ V2.0（体验升级）
4. 每个需求包含：需求ID、标题、描述、优先级、目标版本、关联问题、验收标准
请用 Markdown 表格呈现（表头：| 需求ID | 标题 | 优先级 | 描述 | 关联问题 | 验收标准 |）

### 任务三：测试用例
基于 PRD 生成测试用例。要求：
1. 每个用例**必须能追溯到对应需求（REQ-ID）和原始评论 ID**
2. 包含：用例ID、关联需求、测试标题、前置条件、测试步骤、预期结果、优先级
3. 覆盖正向、边界、异常场景
请用 Markdown 表格呈现（表头：| 用例ID | 关联需求 | 测试标题 | 前置条件 | 测试步骤 | 预期结果 | 优先级 |）

请严格使用上述 Markdown 表格格式，方便后续转成 Word 文档。
"""
    return prompt


# ============================================================================
# 7. Bug 2 修复：用 python-docx 生成真正的 Word 文档
# ============================================================================
def _markdown_table_to_rows(table_lines: list) -> list:
    """把一个 Markdown 表格（多行字符串）拆成二维列表 [[表头...],[数据行...],...]"""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r"^[\s|:\-]+$", line.replace("|", "")):
            continue  # 跳过分隔行 |---|---|
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table_to_doc(doc: Document, table_lines: list):
    """把一组 Markdown 表格行转换成 Word 里的真实表格（add_table）。"""
    rows = _markdown_table_to_rows(table_lines)
    if not rows:
        return

    header = rows[0]
    data_rows = rows[1:]
    n_cols = len(header)

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"  # 'Table Grid' 是 Word 内置样式，带黑色边框，最稳妥

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_data in data_rows:
        cells = table.add_row().cells
        for i in range(n_cols):
            value = row_data[i] if i < len(row_data) else ""
            cells[i].text = value

    doc.add_paragraph("")


def generate_docx_report(analysis_text: str, app_name: str) -> io.BytesIO:
    """把 AI 返回的 Markdown 分析报告转换成真正的 .docx 文件。

    返回：一个 io.BytesIO 对象（里面装着 Word 文件的二进制内容）。
    关键点：用 Document() 创建真实文档、标题用 add_heading、表格用 add_table，
    最后 save 到内存缓冲区 io.BytesIO。
    """
    doc = Document()
    doc.add_heading(f"{app_name} · App Store 评论分析报告", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    sections = re.split(r"\n## ", "\n" + analysis_text)
    for section in sections:
        if not section.strip():
            continue
        lines = section.strip().split("\n")
        first_line = lines[0].strip()

        if first_line.startswith("### "):
            doc.add_heading(first_line[4:], level=2)
        elif first_line.startswith("## "):
            doc.add_heading(first_line[3:], level=1)
        else:
            doc.add_heading(first_line, level=1)

        table_buffer: list = []
        in_table = False
        for line in lines[1:]:
            line = line.rstrip()
            if line.strip().startswith("|") and line.count("|") >= 2:
                table_buffer.append(line)
                in_table = True
            else:
                if in_table and table_buffer:
                    _add_table_to_doc(doc, table_buffer)
                    table_buffer = []
                    in_table = False
                text = line.strip()
                if not text:
                    continue
                if text.startswith(">"):
                    doc.add_paragraph(text[1:].strip(), style="Quote")
                elif text.startswith("- "):
                    doc.add_paragraph(text[2:], style="List Bullet")
                else:
                    doc.add_paragraph(text)
        if in_table and table_buffer:
            _add_table_to_doc(doc, table_buffer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def self_check_docx(buffer: io.BytesIO) -> bool:
    """Bug 2 关键修复之二：生成后立刻用 python-docx 重新打开读一遍，确认文件没坏。"""
    try:
        buffer.seek(0)
        doc = Document(buffer)
        return (len(doc.paragraphs) > 0) or (len(doc.tables) > 0)
    except Exception:
        return False


# ============================================================================
# 8. 零基础友好：python-docx 入门完整示例（逐行中文注释）
# ============================================================================
def docx_beginner_example(output_path: str = "示例文档.docx") -> str:
    """📘 给编程零基础同学的 python-docx 最小可运行示例（每个步骤都有注释）。

    运行方式（终端）：python3 -c "import app; app.docx_beginner_example()"
    """
    from docx import Document
    doc = Document()                                   # 1. 新建空白 Word 文档
    doc.add_heading("我的第一份 Word 文档", level=0)     # 2. 加最大号标题
    doc.add_paragraph("这是用 Python 自动写出来的第一段话。")  # 3. 加一段正文
    doc.add_heading("一个表格示例", level=1)             # 4. 加二级标题
    table = doc.add_table(rows=1, cols=3)              # 5. 建 1行3列 空表
    table.style = "Table Grid"                         # 6. 套带边框样式
    hdr = table.rows[0].cells                          # 7. 取表头行单元格
    hdr[0].text, hdr[1].text, hdr[2].text = "姓名", "年龄", "城市"  # 8. 填表头
    row = table.add_row().cells                        # 9. 新增一行
    row[0].text, row[1].text, row[2].text = "小明", "18", "北京"   # 10. 填数据
    doc.save(output_path)                              # 11. 保存成真实 .docx
    return f"已生成示例文档：{output_path}"


# ============================================================================
# 9. 统一的分析 + 展示流程（供"抓取结果"和"手动粘贴"复用）
# ============================================================================
def analyze_and_display(reviews_data: list, app_name: str, app_id: str, from_cache: bool):
    """对一份评论数据跑 AI 分析，并把结果与下载按钮写入 session_state。

    无论是「RSS 抓取到的评论」还是「用户手动粘贴的评论」，都走这一条同样的流程。
    """
    if not reviews_data:
        st.error("❌ 没有可分析的评论数据")
        return

    total = len(reviews_data)

    # ---- Step 2: 数据概览 ----
    st.divider()
    st.subheader("📊 数据概览")
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        st.metric("评论总数", total)
    with col_b:
        avg_rating = sum(r["rating"] for r in reviews_data) / total if total else 0
        st.metric("平均评分", f"{avg_rating:.1f} ⭐")
    with col_c:
        if from_cache:
            src_label = "📦 缓存"
        elif app_id == "manual":
            src_label = "✍️ 手动粘贴"
        else:
            # 实时抓取：把实际用到的格式（XML / JSON）也标出来，方便验证接口是否生效
            src_label = f"🌐 实时抓取（{LAST_FETCH_SOURCE or '未知'}）"
        st.metric("数据来源", src_label)

    rating_dist = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for r in reviews_data:
        rt = int(r["rating"])
        if rt in rating_dist:
            rating_dist[rt] += 1
    st.caption("评分分布")
    dist_cols = st.columns(5)
    for i, (star, count) in enumerate(rating_dist.items()):
        with dist_cols[i]:
            pct = (count / total * 100) if total else 0
            st.metric(f"{star}⭐", count, f"{pct:.0f}%")
    low_rating = sum(1 for r in reviews_data if r["rating"] <= 2)
    if total and low_rating / total > 0.3:
        st.warning(f"⚠️ 低分评论（1-2星）占比 {low_rating/total*100:.0f}%，用户满意度较低")

    # ---- Step 3: AI 分析 ----
    st.divider()
    st.subheader("🤖 AI 分析中...")
    analysis_progress = st.progress(0, text="准备分析...")
    try:
        analysis_progress.progress(0.1, text="构建分析 Prompt...")
        prompt = build_analysis_prompt(reviews_data, user_goal_global())
        analysis_progress.progress(0.2, text="调用 DeepSeek API（可能需要 30-60 秒）...")
        st.caption(f"分析 {min(total, 200)} 条评论...")
        result_text = call_deepseek(prompt, st.session_state.get("api_key", ""))
        analysis_progress.progress(1.0, text="分析完成！")
        st.session_state.results = {
            "app_name": app_name, "app_id": app_id, "total_reviews": total,
            "avg_rating": avg_rating, "from_cache": from_cache,
            "rating_dist": rating_dist, "analysis": result_text,
        }
    except Exception as e:
        st.error(f"❌ AI 分析失败")
        # 把详细原因折叠显示（避免页面太长），但一眼能看到关键信息
        with st.expander("🔍 查看错误详情", expanded=True):
            st.code(str(e), language="text")
        st.info(
            "💡 排查建议：\n"
            "1）在侧边栏点「🔗 测试连接」验证 Key 是否有效；\n"
            "2）如果提示余额不足，去 https://platform.deepseek.com 充值；\n"
            "3）修复 Key 后重新点「开始分析」会重新实时抓取，无需其他操作。"
        )
        return


def user_goal_global() -> str:
    """取当前界面上的"分析目标"输入框值（手动粘贴流程里也能复用）。"""
    return st.session_state.get("user_goal", "全面分析用户抱怨和痛点")


# ============================================================================
# 10. 主界面
# ============================================================================
def main():
    st.title("📱 App Store Review Analyzer")
    st.markdown("输入美区 App Store 链接，AI 自动分析用户评论 → 生成 PRD → 生成测试用例")

    # ---------- 侧边栏配置 ----------
    with st.sidebar:
        st.header("⚙️ 配置")

        # 把 API Key 存进 session_state，方便 analyze_and_display 复用
        api_key = st.text_input(
            "🔑 DeepSeek API Key",
            type="password",
            placeholder="sk-xxxxxxxxxxxxxxxx",
            help="在 https://platform.deepseek.com 获取",
        )
        st.session_state["api_key"] = api_key

        # 「测试连接」按钮：用最小请求验证 Key 是否有效，避免等到分析时才发现
        if st.button("🔗 测试连接", help="发送一个最小请求到 DeepSeek，验证 Key 是否有效"):
            if not api_key or api_key.strip() == "":
                st.warning("⚠️ 请先填写 API Key 再测试")
            else:
                with st.spinner("正在连接 DeepSeek..."):
                    try:
                        result = call_deepseek("回复一个字：好", api_key, max_tokens=10)
                        st.success(f"✅ 连接成功！DeepSeek 返回：{result.strip()}")
                    except RuntimeError as e:
                        st.error(f"❌ 连接失败\n{e}")

        st.divider()
        st.markdown("### 📊 抓取设置")
        max_reviews = st.slider(
            "最多抓取评论数", min_value=20, max_value=500, value=200, step=10,
            help="抓取太多可能耗时较长",
        )

        st.divider()
        st.markdown("### 📁 缓存状态")
        cache = load_cache()
        if cache:
            st.success(f"✅ 缓存可用 — {cache['total_reviews']} 条评论")
            st.caption(f"App: {cache.get('app_name', 'N/A')}")
            st.caption(f"抓取时间: {cache.get('fetched_at', 'N/A')[:19]}")
        else:
            st.info("📭 暂无缓存")

        if st.button("🗑️ 清除缓存"):
            if CACHE_FILE.exists():
                CACHE_FILE.unlink()
            st.rerun()

        st.divider()
        with st.expander("🧪 链接解析自测（验证 Bug 1 修复）"):
            if st.button("运行自测"):
                for url, expected, actual, passed in run_url_parser_tests():
                    icon = "✅" if passed else "❌"
                    st.write(f"{icon} `{url}`")
                    st.caption(f"　期望: {expected} | 实际: {actual}")

    # ---------- 主区域输入 ----------
    col1, col2 = st.columns([3, 1])
    with col1:
        app_url = st.text_input(
            "🔗 美区 App Store 链接",
            placeholder="https://apps.apple.com/us/app/xxx/id123456789",
        )
    with col2:
        user_goal = st.text_input("🎯 分析目标", placeholder="如：关注订阅问题", value="")
        st.session_state["user_goal"] = user_goal or "全面分析用户抱怨和痛点"

    if not user_goal:
        user_goal = "全面分析用户抱怨和痛点"

    # 🐞 Bug 1 修复点：实时把解析结果显示出来，并做错误检查
    if app_url:
        parsed_id, parse_err = parse_apple_store_url(app_url)
        if parse_err:
            st.error(f"❌ {parse_err}")
        else:
            st.success(f"✅ 链接解析成功 — 美区 App ID：**{parsed_id}**")

    cache_data = load_cache()
    if cache_data and not app_url:
        st.info(
            f"💡 检测到本地缓存（{cache_data['total_reviews']} 条评论，可选离线参考）。"
            f"输入美区链接即可实时抓取并分析；若抓取失败会提示手动粘贴评论。"
        )

    # ---------- 操作按钮 ----------
    btn_col1, btn_col2 = st.columns([1, 1])
    with btn_col1:
        start_btn = st.button("🚀 开始分析", type="primary", use_container_width=True, disabled=not api_key)
    with btn_col2:
        fetch_only_btn = st.button("📥 仅抓取评论", use_container_width=True)

    if not api_key:
        st.warning("⚠️ 请先在侧边栏填写 DeepSeek API Key")

    if "results" not in st.session_state:
        st.session_state.results = None
    if "need_manual" not in st.session_state:
        st.session_state.need_manual = False

    # ========== 仅抓取评论 ==========
    if fetch_only_btn and app_url:
        app_id, parse_err = parse_apple_store_url(app_url)
        if parse_err:
            st.error(f"❌ {parse_err}")
            return
        app_name = extract_app_name(app_url)

        with st.status("正在抓取评论...", expanded=True) as status:
            st.write(f"App: **{app_name}** (ID: {app_id})，使用 **XML + JSON 双通道** 接口")
            try:
                formatted_reviews = fetch_reviews_via_rss(app_id, "us", max_reviews)
                if not formatted_reviews:
                    st.write("RSS 没拿到数据，尝试备用方式...")
                    try:
                        from app_store_scraper import AppStore
                        scraper = AppStore(country="us", app_name=app_name, app_id=app_id)
                        scraper.review(how_many=max_reviews)
                        for r in scraper.reviews:
                            formatted_reviews.append({
                                "id": hashlib.md5(f"{r['userName']}{r['date']}{r['review']}".encode()).hexdigest()[:8],
                                "userName": r.get("userName", "Anonymous"),
                                "title": r.get("title", ""),
                                "review": r.get("review", ""),
                                "rating": r.get("rating", 0),
                                "date": str(r.get("date", "")),
                            })
                    except Exception:
                        pass
                if not formatted_reviews:
                    status.update(label="未获取到评论", state="error")
                    st.error("未获取到任何评论，请检查 App ID 是否正确")
                    return
                save_cache(formatted_reviews, app_name, app_id)
                status.update(label=f"✅ 抓取完成！共 {len(formatted_reviews)} 条评论已缓存", state="complete")
                st.success(f"已保存到 `{CACHE_FILE.name}`")
            except Exception as e:
                status.update(label="抓取失败", state="error")
                st.error(f"抓取出错: {e}")

    # ========== 开始分析 ==========
    if start_btn:
        app_id = None
        app_name = "Unknown App"
        reviews_data = None
        from_cache = False

        if app_url:
            app_id, parse_err = parse_apple_store_url(app_url)
            if parse_err:
                st.error(f"❌ {parse_err}")
                return
            app_name = extract_app_name(app_url)

            progress_bar = st.progress(0)
            st.info(f"🔄 正在从美区 App Store 抓取 **{app_name}** 的评论（XML + JSON 双通道）...")

            try:
                for pct in [0.05, 0.1, 0.15]:
                    time.sleep(0.2)
                    progress_bar.progress(pct, text=f"抓取评论中... {int(pct*100)}%")

                formatted_reviews = fetch_reviews_via_rss(app_id, "us", max_reviews)
                progress_bar.progress(0.5, text="格式化数据...")

                if not formatted_reviews:
                    # v2.0：实时抓取未拿到数据，不再回退缓存，直接进入"手动粘贴"模式（不崩溃）
                    st.session_state.need_manual = True
                    progress_bar.progress(1.0, text="实时抓取无数据")
                else:
                    reviews_data = formatted_reviews
                    save_cache(reviews_data, app_name, app_id)  # 仅作为可选离线产物保存
                    progress_bar.progress(0.6, text="评论已缓存")
            except Exception as e:
                st.warning(f"⚠️ 抓取失败: {e}")
                # v2.0：不再回退缓存，直接进入手动粘贴模式
                st.session_state.need_manual = True
        else:
            # 没填链接 → 没有实时数据可用，直接进入"手动粘贴"模式
            st.session_state.need_manual = True

        # 有数据就走正常分析流程
        if reviews_data:
            analyze_and_display(reviews_data, app_name, app_id, from_cache)

    # ========== 0 条结果时的"手动粘贴"替代入口 ==========
    if st.session_state.need_manual:
        st.divider()
        # 按要求显示的中文提示（明确说明是 Apple 接口限制，不是我们 bug）
        st.error(
            "该 App 当前无法从 Apple 官方获取评论数据"
            "（这是 Apple 接口的已知限制，并非我们产品的 bug），建议：\n"
            "1）稍后重试；\n"
            "2）更换其他热门 App 测试；\n"
            "3）手动粘贴评论文本进行分析"
        )
        st.markdown("### ✍️ 手动粘贴评论（替代数据源）")
        st.caption(
            "格式：每条评论之间空一行；可用「评分: 数字」标注星级（不写默认 3 星）。\n"
            "示例：\n"
            "这个 App 老是闪退，太难受了\n评分: 1\n\n新版本好多了，推荐！\n评分: 5"
        )
        manual_text = st.text_area(
            "把评论粘贴到这里：",
            height=220,
            placeholder="这个 App 老是闪退，太难受了\n评分: 1\n\n新版本好多了，推荐！\n评分: 5",
        )
        if st.button("🚀 用粘贴的评论进行分析"):
            manual_reviews = parse_manual_reviews(manual_text)
            if manual_reviews:
                st.session_state.need_manual = False  # 用完退出手动模式
                st.success(f"✅ 已解析 {len(manual_reviews)} 条手动评论，开始分析...")
                analyze_and_display(manual_reviews, "手动录入的 App", "manual", from_cache=False)
            else:
                st.error("⚠️ 没解析到任何评论，请检查粘贴格式（每条评论之间空一行）")

    # ========== 展示分析结果 ==========
    if st.session_state.results:
        res = st.session_state.results
        st.divider()
        st.header(f"📋 分析结果 — {res['app_name']}")
        st.caption(
            f"App ID: {res['app_id']} | 评论数: {res['total_reviews']} | "
            f"平均评分: {res['avg_rating']:.1f} | 数据源: "
            f"{'缓存' if res['from_cache'] else ('手动粘贴' if res['app_id']=='manual' else '实时')}"
        )
        st.markdown(res["analysis"])

        # ---------- 下载区 ----------
        st.divider()
        st.subheader("📥 下载报告")

        docx_buffer = generate_docx_report(res["analysis"], res["app_name"])
        docx_ok = self_check_docx(docx_buffer)

        if not docx_ok:
            st.error("❌ Word 文档生成后自检失败，已停止提供下载，请检查 python-docx 是否安装正确。")

        dl_col1, dl_col2, dl_col3 = st.columns(3)
        with dl_col1:
            st.download_button(
                "📥 分析报告 (Markdown)",
                data=res["analysis"],
                file_name=f"analysis_{res['app_name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown",
            )
        with dl_col2:
            full_json = json.dumps(res, ensure_ascii=False, indent=2, default=str)
            st.download_button(
                "📥 完整结果 (JSON)",
                data=full_json,
                file_name=f"full_result_{res['app_name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
            )
        with dl_col3:
            if docx_ok:
                st.download_button(
                    "📄 Word 报告（PRD+测试用例）",
                    data=docx_buffer.getvalue(),
                    file_name=f"report_{res['app_name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime=DOCX_MIME,
                )
                st.success("✅ Word 文档已通过自检，可安全下载")
            else:
                st.button("📄 Word 报告（生成失败）", disabled=True)

        with st.expander("📘 python-docx 零基础入门示例（点开看逐行注释代码）"):
            example_src = '''def docx_beginner_example(output_path="示例文档.docx"):
    from docx import Document          # 1. 引入 Document 类（代表一个 Word 文档）
    doc = Document()                   # 2. 新建空白文档
    doc.add_heading("我的第一份 Word 文档", level=0)   # 3. 加最大号标题
    doc.add_paragraph("这是自动写出来的第一段话。")      # 4. 加一段正文
    doc.add_heading("一个表格示例", level=1)           # 5. 加二级标题
    table = doc.add_table(rows=1, cols=3)            # 6. 建 1行3列 空表
    table.style = "Table Grid"                      # 7. 套带边框样式
    hdr = table.rows[0].cells                       # 8. 取表头行单元格
    hdr[0].text, hdr[1].text, hdr[2].text = "姓名","年龄","城市"  # 9. 填表头
    row = table.add_row().cells                     # 10. 新增一行
    row[0].text, row[1].text, row[2].text = "小明","18","北京"    # 11. 填数据
    doc.save(output_path)                           # 12. 保存成真实 .docx 文件'''
            st.code(example_src, language="python")
            st.caption('在终端运行：python3 -c "import app; app.docx_beginner_example()" 即可生成真实 Word 文档。')


if __name__ == "__main__":
    main()
